#!/usr/bin/env python3
# Timestamp: 2025-12-11 16:00:00
# File: /home/ywatanabe/proj/scitex-code/tests/scitex/msword/test_writer.py

"""Tests for scitex.msword.writer module."""

import tempfile
from pathlib import Path

import pytest

# Skip all tests if python-docx is not available
pytestmark = pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)


class TestWordWriterInit:
    """Tests for WordWriter initialization."""

    def test_word_writer_init_with_generic_profile(self):
        """WordWriter should initialize with generic profile."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        assert writer.profile.name == "generic"
        assert writer.template_path is None

    def test_word_writer_init_with_template(self):
        """WordWriter should accept template path."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile, template_path="/some/path.docx")

        assert writer.template_path == "/some/path.docx"


class TestWordWriterWrite:
    """Tests for WordWriter.write method."""

    def test_write_simple_document(self):
        """Should write a simple document with paragraphs."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Introduction"},
                {"type": "paragraph", "text": "This is a test paragraph."},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_output.docx"
            writer.write(doc, output_path)

            assert output_path.exists()
            assert output_path.stat().st_size > 0

    def test_write_document_with_headings(self):
        """Should write document with multiple heading levels."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Section 1"},
                {"type": "heading", "level": 2, "text": "Subsection 1.1"},
                {"type": "paragraph", "text": "Content here."},
                {"type": "heading", "level": 2, "text": "Subsection 1.2"},
                {"type": "paragraph", "text": "More content."},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_headings.docx"
            writer.write(doc, output_path)

            assert output_path.exists()

    def test_write_document_with_table(self):
        """Should write document with a table."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Results"},
                {
                    "type": "table",
                    "rows": [
                        ["Header 1", "Header 2", "Header 3"],
                        ["A", "B", "C"],
                        ["D", "E", "F"],
                    ],
                },
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_table.docx"
            writer.write(doc, output_path)

            assert output_path.exists()

    def test_write_document_with_captions(self):
        """Should write document with figure/table captions."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Figures"},
                {
                    "type": "caption",
                    "caption_type": "figure",
                    "number": 1,
                    "caption_text": "Sample figure caption",
                },
                {
                    "type": "caption",
                    "caption_type": "table",
                    "number": 1,
                    "caption_text": "Sample table caption",
                },
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_captions.docx"
            writer.write(doc, output_path)

            assert output_path.exists()

    def test_write_document_with_formatted_runs(self):
        """Should write document with bold/italic formatting."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {
                    "type": "paragraph",
                    "text": "Mixed formatting",
                    "runs": [
                        {"text": "Normal "},
                        {"text": "bold", "bold": True},
                        {"text": " and "},
                        {"text": "italic", "italic": True},
                        {"text": " text."},
                    ],
                },
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_formatting.docx"
            writer.write(doc, output_path)

            assert output_path.exists()


class TestWordWriterReferences:
    """Tests for reference writing functionality."""

    def test_write_references(self):
        """Should write reference entries."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "References"},
                {
                    "type": "reference-paragraph",
                    "ref_number": 1,
                    "ref_text": "Author A. Title A. Journal 2024.",
                },
                {
                    "type": "reference-paragraph",
                    "ref_number": 2,
                    "ref_text": "Author B. Title B. Journal 2023.",
                },
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_refs.docx"
            writer.write(doc, output_path)

            assert output_path.exists()


class TestWordWriterListItems:
    """Tests for list item writing functionality."""

    def test_write_bullet_list(self):
        """Should write bullet list items."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Bullet Points"},
                {"type": "list-item", "text": "First item", "list_type": "bullet"},
                {"type": "list-item", "text": "Second item", "list_type": "bullet"},
                {"type": "list-item", "text": "Third item", "list_type": "bullet"},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_bullets.docx"
            writer.write(doc, output_path)

            assert output_path.exists()
            assert output_path.stat().st_size > 0

    def test_write_numbered_list(self):
        """Should write numbered list items."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Numbered Steps"},
                {"type": "list-item", "text": "Step one", "list_type": "numbered"},
                {"type": "list-item", "text": "Step two", "list_type": "numbered"},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_numbered.docx"
            writer.write(doc, output_path)

            assert output_path.exists()


class TestWordWriterDoubleAnonymous:
    """Tests for double-anonymous processing."""

    def test_double_anonymous_profile(self):
        """Should apply double-anonymous formatting."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("iop-double-anonymous")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Introduction"},
                {"type": "paragraph", "text": "This is text by John Smith."},
            ],
            "metadata": {
                "author": "John Smith",
            },
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_anon.docx"
            writer.write(doc, output_path)

            assert output_path.exists()


class TestWordWriterStyleExists:
    """Tests for _style_exists method."""

    def test_style_exists_checks_document(self):
        """Should check if style exists in document."""
        import docx

        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = docx.Document()

        # Normal style should exist
        assert writer._style_exists(doc, "Normal") is True

        # Random style should not exist
        assert writer._style_exists(doc, "NonExistentStyleXYZ123") is False


class TestWordWriterEmptyDocument:
    """Tests for handling empty documents."""

    def test_write_empty_blocks(self):
        """Should handle document with no blocks."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_empty.docx"
            writer.write(doc, output_path)

            assert output_path.exists()

    def test_write_blocks_with_empty_text(self):
        """Should skip blocks with empty text."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "paragraph", "text": ""},
                {"type": "heading", "level": 1, "text": ""},
                {"type": "paragraph", "text": "Actual content"},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_empty_text.docx"
            writer.write(doc, output_path)

            assert output_path.exists()


class TestWordWriterProfileSettings:
    """Tests for profile-specific writer settings."""

    def test_write_with_resna_profile(self):
        """Should write with RESNA 2-column profile."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("resna-2025")
        writer = WordWriter(profile=profile)

        assert profile.columns == 2

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "INTRODUCTION"},
                {"type": "paragraph", "text": "Some text."},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_resna.docx"
            writer.write(doc, output_path)

            assert output_path.exists()

    def test_write_with_ieee_profile(self):
        """Should write with IEEE profile."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("ieee")
        writer = WordWriter(profile=profile)

        assert profile.columns == 2

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Introduction"},
                {"type": "paragraph", "text": "Paper content."},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_ieee.docx"
            writer.write(doc, output_path)

            assert output_path.exists()

    def test_write_with_springer_profile(self):
        """Should write with Springer profile."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("springer")
        writer = WordWriter(profile=profile)

        assert profile.columns == 1

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Methods"},
                {"type": "paragraph", "text": "Content here."},
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_springer.docx"
            writer.write(doc, output_path)

            assert output_path.exists()


class TestWordWriterPreExportHooks:
    """Tests for pre-export hooks."""

    def test_pre_export_hooks_called(self):
        """Should call pre-export hooks."""
        from scitex.msword import BaseWordProfile
        from scitex.msword.writer import WordWriter

        hook_called = []

        def my_hook(doc):
            hook_called.append(True)
            return doc

        profile = BaseWordProfile(
            name="test-hooks",
            description="Test",
            pre_export_hooks=[my_hook],
        )
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [{"type": "paragraph", "text": "Test"}],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_hooks.docx"
            writer.write(doc, output_path)

            assert len(hook_called) == 1

    def test_pre_export_hooks_can_modify_doc(self):
        """Should allow hooks to modify document."""
        from scitex.msword import BaseWordProfile
        from scitex.msword.writer import WordWriter

        def add_footer(doc):
            doc["blocks"].append({"type": "paragraph", "text": "Generated by SciTeX"})
            return doc

        profile = BaseWordProfile(
            name="test-modify",
            description="Test",
            pre_export_hooks=[add_footer],
        )
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [{"type": "paragraph", "text": "Content"}],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_modify.docx"
            writer.write(doc, output_path)

            assert output_path.exists()


class TestWordWriterRoundTrip:
    """Tests for round-trip read/write functionality."""

    @pytest.fixture
    def sample_docs_path(self):
        """Path to sample documents."""
        return (
            Path(__file__).parent.parent.parent.parent.parent
            / "docs"
            / "MSWORD_MANUSCTIPS"
        )

    def test_read_modify_write(self, sample_docs_path):
        """Should be able to read a document, modify it, and write it back."""
        from scitex.msword import get_profile
        from scitex.msword.reader import WordReader
        from scitex.msword.writer import WordWriter

        docx_path = sample_docs_path / "RESNA 2025 Scientific Paper Template.docx"
        if not docx_path.exists():
            pytest.skip(f"Sample file not found: {docx_path}")

        # Read
        profile = get_profile("generic")
        reader = WordReader(profile=profile, extract_images=False)
        doc = reader.read(docx_path)

        # Modify
        doc["blocks"].append({"type": "paragraph", "text": "Added by test"})

        # Write
        writer = WordWriter(profile=profile)

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "modified.docx"
            writer.write(doc, output_path)

            assert output_path.exists()
            assert output_path.stat().st_size > 0


class TestWordWriterComplexDocument:
    """Tests for complex document structures."""

    def test_write_complete_manuscript(self):
        """Should write a complete manuscript with all sections."""
        from scitex.msword import get_profile
        from scitex.msword.writer import WordWriter

        profile = get_profile("generic")
        writer = WordWriter(profile=profile)

        doc = {
            "blocks": [
                {"type": "heading", "level": 1, "text": "Abstract"},
                {"type": "paragraph", "text": "This is the abstract."},
                {"type": "heading", "level": 1, "text": "Introduction"},
                {"type": "paragraph", "text": "Background information."},
                {"type": "heading", "level": 1, "text": "Methods"},
                {"type": "heading", "level": 2, "text": "Study Design"},
                {"type": "paragraph", "text": "We conducted..."},
                {"type": "heading", "level": 2, "text": "Data Analysis"},
                {"type": "paragraph", "text": "Statistics were..."},
                {"type": "heading", "level": 1, "text": "Results"},
                {"type": "paragraph", "text": "We found that..."},
                {
                    "type": "caption",
                    "caption_type": "figure",
                    "number": 1,
                    "caption_text": "Results overview",
                },
                {
                    "type": "caption",
                    "caption_type": "table",
                    "number": 1,
                    "caption_text": "Summary statistics",
                },
                {
                    "type": "table",
                    "rows": [
                        ["Variable", "Mean", "SD"],
                        ["Age", "45.2", "12.3"],
                        ["BMI", "25.1", "4.5"],
                    ],
                },
                {"type": "heading", "level": 1, "text": "Discussion"},
                {"type": "paragraph", "text": "Our findings suggest..."},
                {"type": "heading", "level": 1, "text": "Conclusions"},
                {"type": "paragraph", "text": "In conclusion..."},
                {"type": "heading", "level": 1, "text": "References"},
                {
                    "type": "reference-paragraph",
                    "ref_number": 1,
                    "ref_text": "Author A. Title. Journal 2024.",
                },
                {
                    "type": "reference-paragraph",
                    "ref_number": 2,
                    "ref_text": "Author B. Title. Journal 2023.",
                },
            ],
            "metadata": {},
            "images": [],
            "references": [],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "test_complete.docx"
            writer.write(doc, output_path)

            assert output_path.exists()
            assert output_path.stat().st_size > 5000  # Should be a reasonable size


if __name__ == "__main__":
    import os

    import pytest

    pytest.main([os.path.abspath(__file__)])

# --------------------------------------------------------------------------------
# Start of Source Code from: /home/ywatanabe/proj/scitex-code/src/scitex/msword/writer.py
# --------------------------------------------------------------------------------
# #!/usr/bin/env python3
# # -*- coding: utf-8 -*-
# # Timestamp: 2025-12-11 15:15:00
# # File: /home/ywatanabe/proj/scitex-code/src/scitex/msword/writer.py
#
# """
# SciTeX writer document -> DOCX converter.
#
# This module exports SciTeX documents to MS Word .docx files,
# applying journal-specific styles and formatting.
# """
#
# from __future__ import annotations
#
# from pathlib import Path
# from typing import Any, Dict, List, Optional
#
# from .profiles import BaseWordProfile
#
# # Lazy import for python-docx
# try:
#     import docx
#     from docx.document import Document as DocxDocument
#     from docx.shared import Inches, Pt, Cm
#     from docx.enum.text import WD_ALIGN_PARAGRAPH
#     from docx.enum.style import WD_STYLE_TYPE
#
#     DOCX_AVAILABLE = True
#     _DOCX_IMPORT_ERROR = None
# except ImportError as exc:
#     DOCX_AVAILABLE = False
#     _DOCX_IMPORT_ERROR = exc
#
#
# class WordWriter:
#     """
#     Export a SciTeX writer document to a DOCX file.
#
#     This writer handles:
#     - Section headings with proper styles
#     - Paragraphs with formatting
#     - Figure and table captions
#     - References section
#     - Image embedding
#     - Journal-specific template application
#     """
#
#     def __init__(
#         self,
#         profile: BaseWordProfile,
#         template_path: Optional[Path] = None,
#     ):
#         """
#         Parameters
#         ----------
#         profile : BaseWordProfile
#             Mapping from writer structures to Word styles.
#         template_path : Path | None
#             Optional path to a Word template (.dotx/.docx) to use as base.
#         """
#         if not DOCX_AVAILABLE:
#             raise ImportError(
#                 "python-docx is required for scitex.msword.WordWriter. "
#                 "Install it via `pip install python-docx`."
#             ) from _DOCX_IMPORT_ERROR
#         self.profile = profile
#         self.template_path = template_path
#
#     def write(
#         self,
#         writer_doc: Dict[str, Any] | Any,
#         path: Path,
#     ) -> None:
#         """
#         Write a SciTeX writer document to a DOCX file.
#
#         Parameters
#         ----------
#         writer_doc : dict | Any
#             Writer document or intermediate structure.
#         path : Path
#             Output path for the DOCX file.
#         """
#         # Create document (from template if specified)
#         if self.template_path and Path(self.template_path).exists():
#             doc = docx.Document(str(self.template_path))
#             # Clear existing content but keep styles
#             self._clear_document_content(doc)
#         else:
#             doc = docx.Document()
#
#         # Run pre-export hooks
#         for hook in self.profile.pre_export_hooks:
#             writer_doc = hook(writer_doc)
#
#         # Extract blocks from writer_doc
#         if isinstance(writer_doc, dict) and "blocks" in writer_doc:
#             blocks = writer_doc["blocks"]
#             images = writer_doc.get("images", [])
#         else:
#             blocks = list(writer_doc)
#             images = []
#
#         # Build image lookup by hash
#         image_lookup = {img.get("hash"): img for img in images if "hash" in img}
#
#         # Process each block
#         for block in blocks:
#             self._add_block(doc, block, image_lookup)
#
#         # Apply double-anonymous processing if needed
#         if self.profile.double_anonymous:
#             self._apply_double_anonymous(doc, writer_doc)
#
#         # Save document
#         doc.save(str(path))
#
#     def _clear_document_content(self, doc: DocxDocument) -> None:
#         """Clear document content while preserving styles."""
#         for element in doc.element.body[:]:
#             doc.element.body.remove(element)
#
#     def _add_block(
#         self,
#         doc: DocxDocument,
#         block: Dict[str, Any],
#         image_lookup: Dict[str, Any],
#     ) -> None:
#         """Add a single block to the document."""
#         btype = block.get("type", "paragraph")
#         text = block.get("text", "")
#
#         if not text and btype not in ("table", "image"):
#             return
#
#         if btype == "heading":
#             level = block.get("level", 1)
#             self._add_heading(doc, text, level)
#
#         elif btype == "caption":
#             self._add_caption(doc, block)
#
#         elif btype == "reference-paragraph":
#             self._add_reference(doc, block)
#
#         elif btype == "table":
#             self._add_table(doc, block)
#
#         elif btype == "image":
#             self._add_image(doc, block, image_lookup)
#
#         elif btype == "list-item":
#             self._add_list_item(doc, block)
#
#         else:
#             # Default: paragraph
#             self._add_paragraph(doc, text, block.get("runs"))
#
#     def _add_heading(
#         self,
#         doc: DocxDocument,
#         text: str,
#         level: int,
#     ) -> None:
#         """Add a heading paragraph at the given logical level."""
#         style_name = self.profile.heading_styles.get(level)
#
#         if style_name and self._style_exists(doc, style_name):
#             p = doc.add_paragraph(text)
#             p.style = style_name
#         else:
#             # Fallback to built-in heading
#             doc.add_heading(text, level=min(level, 9))
#
#     def _add_paragraph(
#         self,
#         doc: DocxDocument,
#         text: str,
#         runs: Optional[List[Dict[str, Any]]] = None,
#     ) -> None:
#         """Add a paragraph with optional formatted runs."""
#         p = doc.add_paragraph()
#
#         if runs:
#             # Add formatted runs
#             for run_data in runs:
#                 run = p.add_run(run_data.get("text", ""))
#                 if run_data.get("bold"):
#                     run.bold = True
#                 if run_data.get("italic"):
#                     run.italic = True
#                 if run_data.get("underline"):
#                     run.underline = True
#                 if run_data.get("font_size"):
#                     run.font.size = Pt(run_data["font_size"])
#                 if run_data.get("font_name"):
#                     run.font.name = run_data["font_name"]
#         else:
#             p.add_run(text)
#
#         # Apply normal style
#         if self._style_exists(doc, self.profile.normal_style):
#             try:
#                 p.style = self.profile.normal_style
#             except Exception:
#                 pass
#
#     def _add_caption(
#         self,
#         doc: DocxDocument,
#         block: Dict[str, Any],
#     ) -> None:
#         """Add a figure or table caption."""
#         caption_type = block.get("caption_type", "")
#         number = block.get("number", "")
#         caption_text = block.get("caption_text", block.get("text", ""))
#
#         # Build caption text
#         if caption_type == "figure" and number:
#             full_text = f"Figure {number}. {caption_text}"
#         elif caption_type == "table" and number:
#             full_text = f"Table {number}. {caption_text}"
#         else:
#             full_text = block.get("text", caption_text)
#
#         p = doc.add_paragraph(full_text)
#
#         if self._style_exists(doc, self.profile.caption_style):
#             try:
#                 p.style = self.profile.caption_style
#             except Exception:
#                 pass
#
#     def _add_reference(
#         self,
#         doc: DocxDocument,
#         block: Dict[str, Any],
#     ) -> None:
#         """Add a reference entry."""
#         ref_number = block.get("ref_number")
#         ref_text = block.get("ref_text", block.get("text", ""))
#
#         if ref_number is not None:
#             full_text = f"[{ref_number}] {ref_text}"
#         else:
#             full_text = ref_text
#
#         p = doc.add_paragraph(full_text)
#
#         if self._style_exists(doc, self.profile.normal_style):
#             try:
#                 p.style = self.profile.normal_style
#             except Exception:
#                 pass
#
#     def _add_table(
#         self,
#         doc: DocxDocument,
#         block: Dict[str, Any],
#     ) -> None:
#         """Add a table."""
#         rows = block.get("rows", [])
#         if not rows:
#             return
#
#         num_rows = len(rows)
#         num_cols = len(rows[0]) if rows else 0
#
#         table = doc.add_table(rows=num_rows, cols=num_cols)
#         table.style = "Table Grid"
#
#         for i, row_data in enumerate(rows):
#             row = table.rows[i]
#             for j, cell_text in enumerate(row_data):
#                 if j < len(row.cells):
#                     row.cells[j].text = str(cell_text)
#
#     def _add_image(
#         self,
#         doc: DocxDocument,
#         block: Dict[str, Any],
#         image_lookup: Dict[str, Any],
#     ) -> None:
#         """Add an image."""
#         image_hash = block.get("image_hash")
#         image_data = block.get("data")
#
#         if image_hash and image_hash in image_lookup:
#             image_info = image_lookup[image_hash]
#             image_data = image_info.get("data")
#
#         if image_data:
#             from io import BytesIO
#
#             image_stream = BytesIO(image_data)
#             width = block.get("width_inches", 5.0)
#             doc.add_picture(image_stream, width=Inches(width))
#
#     def _add_list_item(
#         self,
#         doc: DocxDocument,
#         block: Dict[str, Any],
#     ) -> None:
#         """Add a list item (bullet or numbered)."""
#         text = block.get("text", "")
#         list_type = block.get("list_type", "bullet")
#
#         p = doc.add_paragraph(text)
#
#         style_key = "bullet" if list_type == "bullet" else "numbered"
#         style_name = self.profile.list_styles.get(style_key)
#
#         if style_name and self._style_exists(doc, style_name):
#             try:
#                 p.style = style_name
#             except Exception:
#                 pass
#
#     def _style_exists(self, doc: DocxDocument, style_name: str) -> bool:
#         """Check if a style exists in the document."""
#         try:
#             _ = doc.styles[style_name]
#             return True
#         except KeyError:
#             return False
#
#     def _apply_double_anonymous(
#         self,
#         doc: DocxDocument,
#         writer_doc: Dict[str, Any],
#     ) -> None:
#         """
#         Apply double-anonymous formatting.
#
#         This removes or masks author-identifying information.
#         """
#         # Get author info to mask
#         metadata = writer_doc.get("metadata", {})
#         author = metadata.get("author", "")
#
#         if not author:
#             return
#
#         # Search and replace author names with placeholder
#         # This is a simple implementation; more sophisticated
#         # masking may be needed for real use
#         for para in doc.paragraphs:
#             if author.lower() in para.text.lower():
#                 for run in para.runs:
#                     if author.lower() in run.text.lower():
#                         # Mask author name
#                         import re
#
#                         run.text = re.sub(
#                             re.escape(author),
#                             "[Author]",
#                             run.text,
#                             flags=re.IGNORECASE,
#                         )
#
#
# __all__ = ["WordWriter"]

# --------------------------------------------------------------------------------
# End of Source Code from: /home/ywatanabe/proj/scitex-code/src/scitex/msword/writer.py
# --------------------------------------------------------------------------------
