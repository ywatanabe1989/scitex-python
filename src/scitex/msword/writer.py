#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2025-12-11 15:15:00
# File: /home/ywatanabe/proj/scitex-code/src/scitex/msword/writer.py

"""
SciTeX writer document -> DOCX converter.

This module exports SciTeX documents to MS Word .docx files,
applying journal-specific styles and formatting.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

from .profiles import BaseWordProfile

# Lazy import for python-docx
try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    DOCX_AVAILABLE = True
    _DOCX_IMPORT_ERROR = None
except ImportError as exc:
    DOCX_AVAILABLE = False
    _DOCX_IMPORT_ERROR = exc


class WordWriter:
    """
    Export a SciTeX writer document to a DOCX file.

    This writer handles:
    - Section headings with proper styles
    - Paragraphs with formatting
    - Figure and table captions
    - References section
    - Image embedding
    - Journal-specific template application
    """

    def __init__(
        self,
        profile: BaseWordProfile,
        template_path: Optional[Path] = None,
    ):
        """
        Parameters
        ----------
        profile : BaseWordProfile
            Mapping from writer structures to Word styles.
        template_path : Path | None
            Optional path to a Word template (.dotx/.docx) to use as base.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for scitex.msword.WordWriter. "
                "Install it via `pip install python-docx`."
            ) from _DOCX_IMPORT_ERROR
        self.profile = profile
        self.template_path = template_path

    def write(
        self,
        writer_doc: Dict[str, Any] | Any,
        path: Path,
    ) -> None:
        """
        Write a SciTeX writer document to a DOCX file.

        Parameters
        ----------
        writer_doc : dict | Any
            Writer document or intermediate structure.
        path : Path
            Output path for the DOCX file.
        """
        # Create document (from template if specified)
        if self.template_path and Path(self.template_path).exists():
            doc = docx.Document(str(self.template_path))
            # Clear existing content but keep styles
            self._clear_document_content(doc)
        else:
            doc = docx.Document()

        # Run pre-export hooks
        for hook in self.profile.pre_export_hooks:
            writer_doc = hook(writer_doc)

        # Extract blocks from writer_doc
        if isinstance(writer_doc, dict) and "blocks" in writer_doc:
            blocks = writer_doc["blocks"]
            images = writer_doc.get("images", [])
        else:
            blocks = list(writer_doc)
            images = []

        # Build image lookup by hash
        image_lookup = {img.get("hash"): img for img in images if "hash" in img}

        # Process each block
        for block in blocks:
            self._add_block(doc, block, image_lookup)

        # Apply double-anonymous processing if needed
        if self.profile.double_anonymous:
            self._apply_double_anonymous(doc, writer_doc)

        # Save document
        doc.save(str(path))

    def _clear_document_content(self, doc: DocxDocument) -> None:
        """Clear document content while preserving styles."""
        for element in doc.element.body[:]:
            doc.element.body.remove(element)

    def _add_block(
        self,
        doc: DocxDocument,
        block: Dict[str, Any],
        image_lookup: Dict[str, Any],
    ) -> None:
        """Add a single block to the document."""
        btype = block.get("type", "paragraph")
        text = block.get("text", "")

        if not text and btype not in ("table", "image"):
            return

        if btype == "heading":
            level = block.get("level", 1)
            self._add_heading(doc, text, level)

        elif btype == "caption":
            self._add_caption(doc, block)

        elif btype == "reference-paragraph":
            self._add_reference(doc, block)

        elif btype == "table":
            self._add_table(doc, block)

        elif btype == "image":
            self._add_image(doc, block, image_lookup)

        elif btype == "list-item":
            self._add_list_item(doc, block)

        else:
            # Default: paragraph
            self._add_paragraph(doc, text, block.get("runs"))

    def _add_heading(
        self,
        doc: DocxDocument,
        text: str,
        level: int,
    ) -> None:
        """Add a heading paragraph at the given logical level."""
        style_name = self.profile.heading_styles.get(level)

        if style_name and self._style_exists(doc, style_name):
            p = doc.add_paragraph(text)
            p.style = style_name
        else:
            # Fallback to built-in heading
            doc.add_heading(text, level=min(level, 9))

    def _add_paragraph(
        self,
        doc: DocxDocument,
        text: str,
        runs: Optional[List[Dict[str, Any]]] = None,
    ) -> None:
        """Add a paragraph with optional formatted runs."""
        p = doc.add_paragraph()

        if runs:
            # Add formatted runs
            for run_data in runs:
                run = p.add_run(run_data.get("text", ""))
                if run_data.get("bold"):
                    run.bold = True
                if run_data.get("italic"):
                    run.italic = True
                if run_data.get("underline"):
                    run.underline = True
                if run_data.get("font_size"):
                    run.font.size = Pt(run_data["font_size"])
                if run_data.get("font_name"):
                    run.font.name = run_data["font_name"]
        else:
            p.add_run(text)

        # Apply normal style
        if self._style_exists(doc, self.profile.normal_style):
            try:
                p.style = self.profile.normal_style
            except Exception:
                pass

    def _add_caption(
        self,
        doc: DocxDocument,
        block: Dict[str, Any],
    ) -> None:
        """Add a figure or table caption."""
        caption_type = block.get("caption_type", "")
        number = block.get("number", "")
        caption_text = block.get("caption_text", block.get("text", ""))

        # Build caption text
        if caption_type == "figure" and number:
            full_text = f"Figure {number}. {caption_text}"
        elif caption_type == "table" and number:
            full_text = f"Table {number}. {caption_text}"
        else:
            full_text = block.get("text", caption_text)

        p = doc.add_paragraph(full_text)

        if self._style_exists(doc, self.profile.caption_style):
            try:
                p.style = self.profile.caption_style
            except Exception:
                pass

    def _add_reference(
        self,
        doc: DocxDocument,
        block: Dict[str, Any],
    ) -> None:
        """Add a reference entry."""
        ref_number = block.get("ref_number")
        ref_text = block.get("ref_text", block.get("text", ""))

        if ref_number is not None:
            full_text = f"[{ref_number}] {ref_text}"
        else:
            full_text = ref_text

        p = doc.add_paragraph(full_text)

        if self._style_exists(doc, self.profile.normal_style):
            try:
                p.style = self.profile.normal_style
            except Exception:
                pass

    def _add_table(
        self,
        doc: DocxDocument,
        block: Dict[str, Any],
    ) -> None:
        """Add a table."""
        rows = block.get("rows", [])
        if not rows:
            return

        num_rows = len(rows)
        num_cols = len(rows[0]) if rows else 0

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = str(cell_text)

    def _add_image(
        self,
        doc: DocxDocument,
        block: Dict[str, Any],
        image_lookup: Dict[str, Any],
    ) -> None:
        """Add an image."""
        image_hash = block.get("image_hash")
        image_data = block.get("data")

        if image_hash and image_hash in image_lookup:
            image_info = image_lookup[image_hash]
            image_data = image_info.get("data")

        if image_data:
            from io import BytesIO

            image_stream = BytesIO(image_data)
            width = block.get("width_inches", 5.0)
            doc.add_picture(image_stream, width=Inches(width))

    def _add_list_item(
        self,
        doc: DocxDocument,
        block: Dict[str, Any],
    ) -> None:
        """Add a list item (bullet or numbered)."""
        text = block.get("text", "")
        list_type = block.get("list_type", "bullet")

        p = doc.add_paragraph(text)

        style_key = "bullet" if list_type == "bullet" else "numbered"
        style_name = self.profile.list_styles.get(style_key)

        if style_name and self._style_exists(doc, style_name):
            try:
                p.style = style_name
            except Exception:
                pass

    def _style_exists(self, doc: DocxDocument, style_name: str) -> bool:
        """Check if a style exists in the document."""
        try:
            _ = doc.styles[style_name]
            return True
        except KeyError:
            return False

    def _apply_double_anonymous(
        self,
        doc: DocxDocument,
        writer_doc: Dict[str, Any],
    ) -> None:
        """
        Apply double-anonymous formatting.

        This removes or masks author-identifying information.
        """
        # Get author info to mask
        metadata = writer_doc.get("metadata", {})
        author = metadata.get("author", "")

        if not author:
            return

        # Search and replace author names with placeholder
        # This is a simple implementation; more sophisticated
        # masking may be needed for real use
        for para in doc.paragraphs:
            if author.lower() in para.text.lower():
                for run in para.runs:
                    if author.lower() in run.text.lower():
                        # Mask author name
                        import re

                        run.text = re.sub(
                            re.escape(author),
                            "[Author]",
                            run.text,
                            flags=re.IGNORECASE,
                        )


__all__ = ["WordWriter"]
